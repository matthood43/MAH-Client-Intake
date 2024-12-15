# dropbox_utils.py
import dropbox
from dropbox.oauth import DropboxOAuth2FlowNoRedirect
from config import (
    DROPBOX_CLIENT_ID,
    DROPBOX_CLIENT_SECRET,
    DROPBOX_TOKEN_FILE
)
from encryption import encrypt_data, decrypt_data
import webbrowser
import logging
import os
import tkinter as tk
from tkinter import simpledialog, messagebox

# Configure logging with rotation
from logging.handlers import RotatingFileHandler

logger = logging.getLogger()
logger.setLevel(logging.INFO)

handler = RotatingFileHandler('../dropbox_utils.log', maxBytes=1000000, backupCount=5)
formatter = logging.Formatter('%(asctime)s:%(levelname)s:%(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

def authenticate(root):
    """
    Performs the OAuth2 authentication flow.
    Guides the user to authorize the app and retrieves the access token via GUI.
    """
    # Initialize the OAuth2 flow without redirect
    auth_flow = DropboxOAuth2FlowNoRedirect(
        DROPBOX_CLIENT_ID,      # Positional argument: app_key
        DROPBOX_CLIENT_SECRET   # Positional argument: app_secret
    )

    authorize_url = auth_flow.start()
    logging.info("Opening browser for Dropbox OAuth2 authentication.")
    print("Opening browser for Dropbox OAuth2 authentication...")
    webbrowser.open(authorize_url)

    # Prompt the user to enter the authorization code via GUI
    auth_code = simpledialog.askstring(
        "Dropbox Authentication",
        f"1. Go to: {authorize_url}\n"
        "2. Click 'Allow' (you might have to log in first).\n"
        "3. Copy the authorization code.\n\n"
        "Enter the authorization code here:",
        parent=root
    )

    if not auth_code:
        messagebox.showerror("Authentication Failed", "No authorization code entered.")
        raise ValueError("No authorization code entered.")

    try:
        oauth_result = auth_flow.finish(auth_code.strip())
    except Exception as e:
        logging.error(f"Error during Dropbox OAuth2 flow: {e}")
        messagebox.showerror("Authentication Error", f"Failed to authenticate with Dropbox:\n{e}")
        raise ValueError("Failed to authenticate with Dropbox. Please try again.")

    access_token = oauth_result.access_token
    # Encrypt and save the access token
    encrypted_token = encrypt_data(access_token)
    with open(DROPBOX_TOKEN_FILE, "w") as token_file:
        token_file.write(encrypted_token)
    logging.info("Dropbox access token obtained and saved securely.")
    messagebox.showinfo("Authentication Successful", "Dropbox has been authenticated successfully.")
    print("Authentication successful. Access token saved.")

    return access_token

def load_access_token():
    """
    Loads the encrypted access token from the token file.
    Returns the decrypted access token or None if not found.
    """
    if not os.path.exists(DROPBOX_TOKEN_FILE):
        return None
    try:
        with open(DROPBOX_TOKEN_FILE, "r") as token_file:
            encrypted_token = token_file.read()
        access_token = decrypt_data(encrypted_token)
        return access_token
    except Exception as e:
        logging.error(f"Error loading access token: {e}")
        return None

def get_dropbox_client(root=None):
    """
    Returns an authenticated Dropbox client.
    If no access token is found, initiates the OAuth2 flow.
    The 'root' parameter is the Tkinter root window, required for GUI dialogs.
    """
    access_token = load_access_token()
    if not access_token:
        logging.info("No access token found. Initiating OAuth2 flow.")
        if root:
            access_token = authenticate(root)
        else:
            # Create a temporary hidden root window for dialogs
            temp_root = tk.Tk()
            temp_root.withdraw()
            access_token = authenticate(temp_root)
            temp_root.destroy()

    db_client = dropbox.Dropbox(access_token)
    try:
        db_client.users_get_current_account()
        logging.info("Dropbox client authenticated successfully.")
    except dropbox.exceptions.AuthError as e:
        logging.error(f"Invalid access token: {e}")
        if root:
            messagebox.showerror("Authentication Error", "Invalid Dropbox access token. Please re-authenticate.")
        raise ValueError("Invalid Dropbox access token. Please re-authenticate.")

    return db_client

def upload_to_dropbox(local_path, dropbox_path, root=None):
    """
    Uploads a file to Dropbox.
    """
    try:
        db_client = get_dropbox_client(root)
        with open(local_path, "rb") as f:
            db_client.files_upload(f.read(), dropbox_path, mode=dropbox.files.WriteMode.overwrite)
        logging.info(f"Uploaded {local_path} to Dropbox at {dropbox_path}.")
    except Exception as e:
        logging.error(f"Error uploading to Dropbox: {e}")
        if root:
            messagebox.showerror("Upload Error", f"Failed to upload {local_path} to Dropbox:\n{e}")
        raise

def create_dropbox_folders(client_name, root=None):
    """
    Creates necessary folders in Dropbox for the client.
    """
    try:
        db_client = get_dropbox_client(root)
        folders = [
            f"/{client_name}",
            f"/{client_name}/Correspondence",
            f"/{client_name}/Documents",
            f"/{client_name}/Invoices"
        ]
        for folder in folders:
            try:
                db_client.files_create_folder_v2(folder)
                logging.info(f"Created Dropbox folder: {folder}")
            except dropbox.exceptions.ApiError as e:
                if e.error.is_path() and e.error.get_path().is_conflict():
                    logging.info(f"Dropbox folder already exists: {folder}")
                else:
                    logging.error(f"Error creating folder {folder}: {e}")
                    if root:
                        messagebox.showerror("Folder Creation Error", f"Failed to create folder {folder}:\n{e}")
                    raise
    except Exception as e:
        logging.error(f"Error in create_dropbox_folders: {e}")
        raise

def process_fee_agreement(client_data, output_path):
    """
    Processes the fee agreement document.
    """
    # Example implementation using python-docx
    from docx import Document

    document = Document()

    document.add_heading('Fee Agreement', 0)

    for key, value in client_data.items():
        if key != "EncryptedData":  # Exclude encrypted data field
            # Replace underscores with spaces and capitalize words
            formatted_key = ' '.join(word.capitalize() for word in key.split('_'))
            document.add_paragraph(f"{formatted_key}: {value}")

    document.save(output_path)
    logging.info(f"Fee agreement document created at {output_path}.")
