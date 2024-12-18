import os
import json
import pyodbc
import dropbox
import base64
import re
import requests
import logging
from cryptography.fernet import Fernet
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
from pathlib import Path

# PyQt5 imports for GUI
import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QPushButton,
    QFormLayout, QMessageBox, QCheckBox, QComboBox
)
from PyQt5.QtCore import QObject, pyqtSignal, QThread
from http.server import HTTPServer, BaseHTTPRequestHandler
import threading
import webbrowser
from urllib.parse import urlparse, parse_qs

# ------------------------------ Configuration ------------------------------ #

# Configure logging
logging.basicConfig(
    filename='client_intake.log',
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.DEBUG  # Set to DEBUG for comprehensive logging
)
console = logging.StreamHandler()
console.setLevel(logging.INFO)  # Adjust as needed
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console.setFormatter(formatter)
logging.getLogger('').addHandler(console)

# Load environment variables from .env file
load_dotenv()

# Dropbox OAuth2 Credentials from .env
DROPBOX_CLIENT_ID = os.getenv("DROPBOX_CLIENT_ID")
DROPBOX_CLIENT_SECRET = os.getenv("DROPBOX_CLIENT_SECRET")
DROPBOX_REDIRECT_URI = os.getenv("DROPBOX_REDIRECT_URI")  # e.g., "http://localhost:5000/"

# Validate Dropbox credentials
if not DROPBOX_CLIENT_ID or not DROPBOX_CLIENT_SECRET or not DROPBOX_REDIRECT_URI:
    logging.critical("Dropbox CLIENT_ID, CLIENT_SECRET, and REDIRECT_URI must be set in the .env file.")
    raise ValueError("Dropbox CLIENT_ID, CLIENT_SECRET, and REDIRECT_URI must be set in the .env file.")

# Harvest Credentials (Secure these as well)
HARVEST_ACCOUNT_ID = os.getenv("HARVEST_ACCOUNT_ID")
HARVEST_TOKEN = os.getenv("HARVEST_TOKEN")  # Replace with actual token

# Bosch Client ID in Harvest (replace with the actual client_id for Bosch)
HARVEST_BOSCH_CLIENT_ID = int(os.getenv("HARVEST_BOSCH_CLIENT_ID", 0))

# Database Path (Ensure this path is correct and accessible)
LOCAL_DB_PATH = os.getenv("LOCAL_DB_PATH", default=r"C:\Users\Matt\Dropbox\Client Intake Program\Clients1.accdb")

# Fee Agreement Template Path
FEE_AGREEMENT_TEMPLATE_PATH = os.getenv("FEE_AGREEMENT_TEMPLATE_PATH",
                                        default=r"C:\Users\Matt\Dropbox\Client Intake Program\Fee_Agreement.docx")

# ------------------------------ Encryption Setup ------------------------------ #

def generate_key():
    """
    Generates a new encryption key and saves it to 'encryption_key.key'.
    """
    key = Fernet.generate_key()
    with open("encryption_key.key", "wb") as key_file:
        key_file.write(key)
    logging.info("Encryption key generated and saved to 'encryption_key.key'.")
    return key

def load_key():
    """
    Loads the encryption key from 'encryption_key.key'.
    If the key file does not exist, it generates a new key.
    """
    key_path = Path("encryption_key.key")
    if not key_path.exists():
        logging.info("Encryption key not found. Generating a new key.")
        generate_key()
    with key_path.open("rb") as key_file:
        key = key_file.read()
    return key

def encrypt_data(data):
    """
    Encrypts the provided data using Fernet symmetric encryption.

    Args:
        data (str or None): The plaintext data to encrypt.

    Returns:
        str: The encrypted data encoded in base64.
    """
    try:
        if data is None:
            data = ""
            logging.debug("Received None, converting to empty string for encryption.")
        key = load_key()
        f = Fernet(key)
        encrypted_data = f.encrypt(data.encode())
        encrypted_b64 = base64.b64encode(encrypted_data).decode('utf-8')
        logging.debug("Data encrypted.")
        return encrypted_b64
    except Exception as e:
        logging.error(f"Encryption failed: {e}", exc_info=True)
        raise

def decrypt_data(encrypted_data):
    """
    Decrypts the provided encrypted data using Fernet symmetric encryption.

    Args:
        encrypted_data (str): The base64-encoded encrypted data.

    Returns:
        str: The decrypted plaintext data.
    """
    try:
        key = load_key()
        f = Fernet(key)
        decrypted_data = f.decrypt(base64.b64decode(encrypted_data))
        decrypted_str = decrypted_data.decode('utf-8')
        logging.debug("Data decrypted.")
        return decrypted_str
    except Exception as e:
        logging.error(f"Decryption failed: {e}", exc_info=True)
        raise

# ------------------------------ Validation Functions ------------------------------ #

def is_valid_email(email):
    """
    Validates the email format using regex.

    Args:
        email (str): The email address to validate.

    Returns:
        bool: True if valid, False otherwise.
    """
    email_regex = r"(^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$)"
    valid = re.match(email_regex, email) is not None
    logging.debug(f"Email validation for '{email}': {valid}")
    return valid

def is_valid_phone(phone):
    """
    Validates the phone number format using regex.

    Args:
        phone (str): The phone number to validate.

    Returns:
        bool: True if valid, False otherwise.
    """
    phone_regex = r"^\+?1?\d{10,15}$"  # Simple regex for phone numbers
    # Remove common formatting characters
    cleaned_phone = re.sub(r"[()\-\s+]", "", phone)
    valid = re.match(phone_regex, cleaned_phone) is not None
    logging.debug(f"Phone validation for '{phone}': {valid}")
    return valid

def is_valid_address(address):
    """
    Validates the address format using regex.

    Args:
        address (str): The address to validate.

    Returns:
        bool: True if valid, False otherwise.
    """
    # Simple regex to check for street, city, state, and zip code
    # Example: "123 Main St, Springfield, IL 62704"
    address_regex = r"^[\d\w\s,.]+,\s*[\w\s]+,\s*[A-Z]{2}\s*\d{5}(-\d{4})?$"
    valid = re.match(address_regex, address) is not None
    logging.debug(f"Address validation for '{address}': {valid}")
    return valid

# ------------------------------ OAuth2 Handler ------------------------------ #

class RedirectHandler(BaseHTTPRequestHandler):
    """
    HTTP request handler to capture the OAuth2 redirect with the authorization code.
    """

    def do_GET(self):
        parsed_url = urlparse(self.path)
        params = parse_qs(parsed_url.query)
        if 'code' in params:
            self.server.auth_code = params['code'][0]
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            self.wfile.write(b"<html><body><h1>Authentication Successful</h1>You can close this window.</body></html>")
            logging.info("Authorization code received.")
            # Signal that auth code has been received
            self.server.auth_event.set()
        else:
            self.send_response(400)
            self.end_headers()
            logging.warning("Authorization code not found in the redirect URI.")

class OAuthHandlerWithServer(QObject):
    """
    Handles the OAuth2 authentication flow with Dropbox using a local HTTP server.
    """
    # Define the signals
    token_received = pyqtSignal(str)
    auth_failed = pyqtSignal(str)

    def __init__(self, client_id, client_secret, redirect_uri):
        super().__init__()
        self.client_id = client_id
        self.client_secret = client_secret
        self.redirect_uri = redirect_uri
        self.server = None
        self.auth_event = threading.Event()

    def start_auth_flow(self):
        # Parse the redirect URI to extract the port
        parsed = urlparse(self.redirect_uri)
        port = parsed.port
        if not port:
            error_msg = "Redirect URI must include a port number."
            logging.error(error_msg)
            self.auth_failed.emit(error_msg)
            return

        handler = RedirectHandler
        try:
            self.server = HTTPServer(('localhost', port), handler)
            self.server.auth_event = self.auth_event  # Pass the event to the server
            server_thread = threading.Thread(target=self.server.serve_forever)
            server_thread.setDaemon(True)
            server_thread.start()
            logging.info(f"Started local HTTP server on port {port} to capture OAuth2 redirect.")
        except Exception as e:
            error_msg = f"Failed to start local server: {e}"
            logging.error(error_msg)
            self.auth_failed.emit(error_msg)
            return

        # Generate the authorization URL
        auth_url = (
            f"https://www.dropbox.com/oauth2/authorize?"
            f"response_type=code&client_id={self.client_id}&redirect_uri={self.redirect_uri}"
        )

        logging.info("Opening web browser for Dropbox OAuth2 authorization.")
        # Open the authorization URL in the default web browser
        try:
            webbrowser.open(auth_url)
            logging.info("Opened web browser for Dropbox OAuth2 authorization.")
        except Exception as e:
            logging.error(f"Failed to open web browser: {e}")
            self.auth_failed.emit(f"Failed to open web browser: {e}")

        # Start a thread to wait for the auth code
        wait_thread = threading.Thread(target=self.wait_for_code)
        wait_thread.start()

    def wait_for_code(self):
        try:
            # Wait until the server sets the auth_code or timeout after 120 seconds
            if not self.auth_event.wait(timeout=120):
                error_msg = "Authentication timed out."
                logging.error(error_msg)
                self.auth_failed.emit(error_msg)
                self.server.shutdown()
                return

            if hasattr(self.server, 'auth_code'):
                code = self.server.auth_code
                logging.info("Exchanging authorization code for access token.")
                # Exchange the authorization code for an access token
                token_url = "https://api.dropboxapi.com/oauth2/token"
                data = {
                    'code': code,
                    'grant_type': 'authorization_code',
                    'client_id': self.client_id,
                    'client_secret': self.client_secret,
                    'redirect_uri': self.redirect_uri
                }
                response = requests.post(token_url, data=data)
                if response.status_code == 200:
                    token_info = response.json()
                    access_token = token_info.get('access_token')
                    if access_token:
                        logging.info("Access token obtained successfully.")
                        self.token_received.emit(access_token)
                    else:
                        error_msg = "Access token not found in response."
                        logging.error(error_msg)
                        self.auth_failed.emit(error_msg)
                else:
                    error_msg = f"Token exchange failed: {response.text}"
                    logging.error(error_msg)
                    self.auth_failed.emit(error_msg)
            else:
                error_msg = "Authorization code not found."
                logging.error(error_msg)
                self.auth_failed.emit(error_msg)
        except Exception as e:
            error_msg = f"An error occurred during token exchange: {e}"
            logging.error(error_msg, exc_info=True)
            self.auth_failed.emit(error_msg)
            self.server.shutdown()

# ------------------------------ Dropbox Client Initializer ------------------------------ #

class DropboxClientInitializer(QObject):
    """
    Initializes the Dropbox client by handling OAuth2 authentication and loading the access token.
    """
    # Define the signals
    token_received = pyqtSignal(str)
    auth_failed = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.db_client = None

    def initialize_dropbox_client(self):
        token_file = Path("dropbox_token.json")
        key_file = Path("encryption_key.key")

        # Check if the token file exists
        if token_file.exists():
            try:
                with key_file.open("rb") as f:
                    key = f.read()
                fernet = Fernet(key)
                with token_file.open("rb") as f:
                    encrypted_data = f.read()
                decrypted_data = fernet.decrypt(encrypted_data)
                tokens = json.loads(decrypted_data.decode('utf-8'))
                access_token = tokens.get("access_token")
                if access_token:
                    logging.info("Access token loaded from 'dropbox_token.json'.")
                    self.db_client = dropbox.Dropbox(access_token)
                    # Verify the token by making a simple API call
                    try:
                        self.db_client.users_get_current_account()
                        logging.info("Dropbox access token is valid.")
                        # Emit signal to indicate successful initialization
                        self.token_received.emit(access_token)
                        return
                    except Exception as e:
                        logging.error(f"Invalid access token: {e}")
                        # Proceed to re-authenticate
                else:
                    logging.warning("Access token not found in decrypted data.")
            except Exception as e:
                logging.error(f"Failed to decrypt or load token file: {e}")

        # If token doesn't exist or decryption failed, initiate OAuth2 flow
        logging.info("Initiating OAuth2 authentication flow for Dropbox.")
        self.oauth_handler = OAuthHandlerWithServer(DROPBOX_CLIENT_ID, DROPBOX_CLIENT_SECRET, DROPBOX_REDIRECT_URI)

        # Connect OAuthHandlerWithServer signals to DropboxClientInitializer signals
        self.oauth_handler.token_received.connect(self.token_received.emit)
        self.oauth_handler.auth_failed.connect(self.auth_failed.emit)

        # Start the OAuth2 flow
        self.oauth_handler.start_auth_flow()

# ------------------------------ Database and Dropbox Functions ------------------------------ #

def upload_to_dropbox(db_client, local_path, dropbox_path):
    """
    Uploads a local file to Dropbox.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        local_path (Path): Path to the local file.
        dropbox_path (str): Destination path in Dropbox.

    Raises:
        Exception: If the upload fails.
    """
    try:
        with local_path.open("rb") as file:
            db_client.files_upload(file.read(), dropbox_path, mode=dropbox.files.WriteMode.overwrite)
            logging.info(f"Uploaded to Dropbox: {dropbox_path}")
    except Exception as e:
        logging.error(f"Error uploading to Dropbox: {e}")
        raise

def download_from_dropbox(db_client, dropbox_path, local_path):
    """
    Downloads a file from Dropbox to a local path.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        dropbox_path (str): Path to the file in Dropbox.
        local_path (Path): Destination path on the local machine.

    Raises:
        Exception: If the download fails.
    """
    try:
        metadata, res = db_client.files_download(path=dropbox_path)
        with local_path.open("wb") as f:
            f.write(res.content)
        logging.info(f"Downloaded from Dropbox: {dropbox_path} to {local_path}")
    except dropbox.exceptions.ApiError as e:
        logging.error(f"Error downloading from Dropbox: {e}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error downloading from Dropbox: {e}")
        raise

def initialize_database(db_client, local_db_path):
    """
    Ensures that the local database exists and is synchronized with Dropbox.
    Establishes a connection to the Microsoft Access database.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        local_db_path (Path): Path to the local Access database.

    Returns:
        pyodbc.Connection: Active database connection.

    Raises:
        FileNotFoundError: If the local database file is missing and cannot be downloaded.
        Exception: If the database cannot be connected.
    """
    db_dropbox_path = "/Client Intake Program/Clients1.accdb"

    # Check if the local database file exists
    if not local_db_path.exists():
        logging.warning(f"Local database file not found at '{local_db_path}'. Attempting to download from Dropbox.")
        try:
            # Attempt to download the database from Dropbox
            download_from_dropbox(db_client, db_dropbox_path, local_db_path)
            logging.info(f"Database downloaded successfully from Dropbox to '{local_db_path}'.")
        except Exception as e:
            logging.error(f"Failed to download database from Dropbox: {e}")
            raise FileNotFoundError(f"Local database file not found and failed to download from Dropbox: {e}")

    try:
        connection_string = (
            r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=' + str(local_db_path)
        )
        # Include password only if applicable
        DATABASE_PASSWORD = os.getenv("DATABASE_PASSWORD")
        if DATABASE_PASSWORD:
            connection_string += r';PWD=' + DATABASE_PASSWORD

        logging.debug(f"Connection string: {connection_string}")
        conn = pyodbc.connect(connection_string)
        logging.info("Database connection established successfully.")
        return conn
    except pyodbc.Error as e:
        logging.error(f"Failed to connect to the database: {e}")
        raise

def list_table_columns(conn, table_name):
    """
    Lists the column names of a specified table in the database.

    Args:
        conn (pyodbc.Connection): Active database connection.
        table_name (str): Name of the table to inspect.

    Returns:
        list: List of column names.

    Raises:
        Exception: If unable to retrieve column information.
    """
    try:
        cursor = conn.cursor()
        # Use the system tables to get column names
        cursor.execute(f"SELECT * FROM [{table_name}] WHERE 1=0")
        columns = [column[0] for column in cursor.description]
        logging.info(f"Columns in '{table_name}' table: {columns}")
        return columns
    except Exception as e:
        logging.error(f"Error retrieving columns for table '{table_name}': {e}")
        raise

def insert_client_data(conn, client_data, encrypted_data):
    """
    Inserts client data into the database.

    Args:
        conn (pyodbc.Connection): Active database connection.
        client_data (dict): Dictionary containing client information.
        encrypted_data (str): Encrypted JSON string of client_data.

    Raises:
        Exception: If the insertion fails.
    """
    try:
        cursor = conn.cursor()
        cursor.execute(
            '''INSERT INTO Clients (
                Client_FirstName, Client_LastName, Opposing_Party, Client_Email, Client_Address, Client_Phone,
                Fee_Arrangement, Contingency_Amount, Hourly_Amount, Primary_Attorney, EncryptedData, IsBosch
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                client_data.get('Client_FirstName'),
                client_data.get('Client_LastName'),
                client_data.get('Opposing_Party'),
                client_data.get('Client_Email'),
                client_data.get('Client_Address'),
                client_data.get('Client_Phone'),
                client_data.get('Fee_Arrangement'),
                client_data.get('Contingency_Amount'),
                client_data.get('Hourly_Amount'),
                client_data.get('Primary_Attorney'),
                encrypted_data,
                client_data.get('IsBosch')
            )
        )
        conn.commit()
        logging.info("Client data stored successfully in the database.")
    except Exception as e:
        logging.error(f"Error inserting data into the database: {e}")
        raise

def create_dropbox_folders(db_client, client_name):
    """
    Creates necessary folders for a new client in Dropbox.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        client_name (str): Full name of the client (e.g., 'John_Doe').

    Raises:
        Exception: If folder creation fails for reasons other than existing folders.
    """
    try:
        main_folder_path = f"/{client_name}"
        subfolders = ["Pleadings", "Correspondence", "Prelitigation", "Case Notes"]

        # Create the main client folder
        try:
            db_client.files_create_folder_v2(main_folder_path)
            logging.info(f"Main folder created: {main_folder_path}")
        except dropbox.exceptions.ApiError as e:
            if e.error.is_path() and e.error.get_path().is_conflict():
                logging.info(f"Main folder already exists: {main_folder_path}")
            else:
                logging.error(f"Error creating main folder '{main_folder_path}': {e}")
                raise

        # Create subfolders within the main client folder
        for subfolder in subfolders:
            subfolder_path = f"{main_folder_path}/{subfolder}"
            try:
                db_client.files_create_folder_v2(subfolder_path)
                logging.info(f"Subfolder created: {subfolder_path}")
            except dropbox.exceptions.ApiError as e:
                if e.error.is_path() and e.error.get_path().is_conflict():
                    logging.info(f"Subfolder already exists: {subfolder_path}")
                else:
                    logging.error(f"Error creating subfolder '{subfolder_path}': {e}")
                    raise
    except Exception as e:
        logging.error(f"Failed to create Dropbox folders for '{client_name}': {e}")
        raise

def create_fee_agreement(client_data, template_path, output_filename):
    """
    Creates a personalized fee agreement document from a local template.

    Args:
        client_data (dict): Dictionary containing client information.
        template_path (str): Local path to the Fee_Agreement.docx template.
        output_filename (str): Desired name for the personalized document (e.g., 'Fee Agreement - John Doe.docx').

    Returns:
        BytesIO: In-memory binary stream of the personalized document.

    Raises:
        FileNotFoundError: If the template file does not exist.
        Exception: If the document creation fails.
    """
    try:
        # Convert template_path to a Path object for better path handling
        template = Path(template_path)
        logging.debug(f"Attempting to load template from: {template}")

        if not template.exists():
            error_msg = f"Fee agreement template not found at '{template}'. Please ensure the file exists."
            logging.error(error_msg)
            raise FileNotFoundError(error_msg)

        # Load the template document
        document = Document(str(template))
        logging.info(f"Loaded template from '{template}'.")

        # Replace placeholders in paragraphs
        for paragraph in document.paragraphs:
            for key, value in client_data.items():
                if value is None:
                    value = ""
                placeholder = f"{{{key}}}"  # e.g., {Client_FirstName}
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    logging.debug(f"Replaced '{placeholder}' with '{value}' in paragraph.")

        # Replace placeholders in tables (if any)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in client_data.items():
                        if value is None:
                            value = ""
                        placeholder = f"{{{key}}}"  # e.g., {Client_FirstName}
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            logging.debug(f"Replaced '{placeholder}' with '{value}' in table cell.")

        # Save the personalized document to a BytesIO stream
        output_stream = BytesIO()
        document.save(output_stream)
        output_stream.seek(0)
        logging.info(f"Personalized fee agreement '{output_filename}' created in memory.")

        return output_stream
    except FileNotFoundError as fnf_error:
        logging.error(fnf_error)
        raise
    except Exception as e:
        logging.error(f"Error creating fee agreement: {e}", exc_info=True)
        raise

def upload_fee_agreement(db_client, fee_agreement_stream, fee_agreement_dropbox_path):
    """
    Uploads the personalized fee agreement to Dropbox.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        fee_agreement_stream (BytesIO): In-memory binary stream of the personalized document.
        fee_agreement_dropbox_path (str): Dropbox path where the document will be uploaded.

    Raises:
        Exception: If the upload fails.
    """
    try:
        db_client.files_upload(fee_agreement_stream.read(), fee_agreement_dropbox_path,
                               mode=dropbox.files.WriteMode.overwrite)
        logging.info(f"Uploaded fee agreement to '{fee_agreement_dropbox_path}'.")
    except Exception as e:
        logging.error(f"Failed to upload fee agreement: {e}", exc_info=True)
        raise

def create_harvest_project(client_data, harvest_bosch_client_id):
    """
    Creates a Harvest project for the client if applicable.

    Args:
        client_data (dict): Dictionary containing client information.
        harvest_bosch_client_id (int): Bosch client ID in Harvest.

    Raises:
        Exception: If project creation fails.
    """
    headers = {
        "User-Agent": "ClientIntakeApp (support@example.com)",
        "Authorization": f"Bearer {HARVEST_TOKEN}",
        "Harvest-Account-Id": HARVEST_ACCOUNT_ID,
        "Content-Type": "application/json"
    }

    project_data = {
        "client_id": harvest_bosch_client_id,
        "name": f"{client_data['Client_FirstName']} {client_data['Client_LastName']}",
        "is_billable": True,
        "bill_by": "Project",
        "hourly_rate": 450.0
    }

    try:
        response = requests.post("https://api.harvestapp.com/v2/projects", json=project_data, headers=headers)
        if response.status_code == 201:
            logging.info("Harvest project created successfully with an hourly rate of 450.")
        else:
            error_msg = f"Failed to create Harvest project: {response.status_code}, {response.text}"
            logging.error(error_msg)
            raise Exception("Harvest project creation failed.")
    except Exception as e:
        logging.error(f"An error occurred while creating Harvest project: {e}")
        raise

def process_fee_agreement(db_client, client_data, client_name):
    """
    Creates and uploads the fee agreement document for the client.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        client_data (dict): Dictionary containing client information.
        client_name (str): Full name of the client (e.g., 'John_Doe').

    Raises:
        Exception: If processing fails.
    """
    try:
        # Define template path and output filename
        template_path = FEE_AGREEMENT_TEMPLATE_PATH
        output_filename = f"Fee Agreement - {client_data['Client_FirstName']} {client_data['Client_LastName']}.docx"

        # Create the personalized fee agreement document
        fee_agreement_stream = create_fee_agreement(client_data, template_path, output_filename)

        # Define the Dropbox path for the fee agreement
        fee_agreement_dropbox_path = f"/{client_name}/Correspondence/{output_filename}"

        # Upload the fee agreement to Dropbox
        upload_fee_agreement(db_client, fee_agreement_stream, fee_agreement_dropbox_path)
    except Exception as e:
        logging.error(f"Failed to process fee agreement: {e}", exc_info=True)
        raise

def process_client_data(db_client, db_conn, client_data):
    """
    Processes the intake data for a new client.

    Args:
        db_client (dropbox.Dropbox): Authenticated Dropbox client.
        db_conn (pyodbc.Connection): Active database connection.
        client_data (dict): Dictionary containing client information.

    Returns:
        str: Client's full name used for folder creation.

    Raises:
        ValueError: If validation fails.
        Exception: If processing fails.
    """
    # Validate email
    if not is_valid_email(client_data['Client_Email']):
        error_msg = "Invalid email address."
        logging.warning(error_msg)
        raise ValueError(error_msg)

    # Validate phone number
    if not is_valid_phone(client_data['Client_Phone']):
        error_msg = "Invalid phone number. Please enter a valid phone number with 10-15 digits."
        logging.warning(error_msg)
        raise ValueError(error_msg)

    # Validate address
    if not is_valid_address(client_data['Client_Address']):
        error_msg = "Invalid address format. Please include street, city, state, and zip code."
        logging.warning(error_msg)
        raise ValueError(error_msg)

    # Optional: Verify table columns
    table_columns = list_table_columns(db_conn, "Clients")
    expected_columns = [
        "Client_FirstName", "Client_LastName", "Opposing_Party", "Client_Email",
        "Client_Address", "Client_Phone", "Fee_Arrangement", "Contingency_Amount",
        "Hourly_Amount", "Primary_Attorney", "EncryptedData", "IsBosch"
    ]
    missing_columns = [col for col in expected_columns if col not in table_columns]
    if missing_columns:
        error_msg = f"The following required columns are missing in the 'Clients' table: {missing_columns}"
        logging.error(error_msg)
        raise ValueError(error_msg)

    # Prepare data for insertion
    # Convert string amounts to floats only if applicable
    try:
        if client_data['Fee_Arrangement'] == "Contingency":
            contingency_percentage = client_data.get('Contingency_Percentage', "0.0")
            if contingency_percentage is None or contingency_percentage == "":
                contingency_percentage = "0.0"
            client_data['Contingency_Amount'] = float(contingency_percentage)
        else:
            client_data['Contingency_Amount'] = 0.0

        if client_data['Fee_Arrangement'] == "Hourly":
            hourly_rate = client_data.get('Hourly_Rate', "0.0")
            if hourly_rate is None or hourly_rate == "":
                hourly_rate = "0.0"
            client_data['Hourly_Amount'] = float(hourly_rate)
        else:
            client_data['Hourly_Amount'] = 0.0
    except ValueError:
        error_msg = "Contingency amount and Hourly amount must be valid numbers."
        logging.error(error_msg)
        raise ValueError(error_msg)

    # Convert 'IsBosch' to Boolean
    client_data['IsBosch'] = client_data.get('IsBosch') == "True"

    # Encrypt all client data and store in 'EncryptedData'
    raw_json = json.dumps(client_data)
    encrypted_data = encrypt_data(raw_json)

    # Insert data into the database
    try:
        insert_client_data(db_conn, client_data, encrypted_data)
    except Exception as e:
        logging.error(f"Inserting client data failed: {e}")
        raise

    # Create Dropbox folders
    client_name = f"{client_data['Client_FirstName']}_{client_data['Client_LastName']}"
    try:
        create_dropbox_folders(db_client, client_name)
    except Exception as e:
        logging.error(f"Creating Dropbox folders failed: {e}")
        raise

    # Create and upload fee agreement document
    try:
        process_fee_agreement(db_client, client_data, client_name)
    except Exception as e:
        logging.error(f"Processing fee agreement failed: {e}")
        raise

    # Create Harvest project if applicable
    if client_data['IsBosch']:
        try:
            create_harvest_project(client_data, HARVEST_BOSCH_CLIENT_ID)
        except Exception as e:
            logging.error(f"Creating Harvest project failed: {e}")
            raise

    return client_name

# ------------------------------ Worker Class for Initialization ------------------------------ #

class InitializationWorker(QObject):
    """
    Worker class to handle Dropbox authentication in a separate thread.
    """
    initialization_complete = pyqtSignal()
    initialization_failed = pyqtSignal(str)

    def __init__(self, db_initializer):
        super().__init__()
        self.db_initializer = db_initializer

    def run(self):
        """
        Executes the Dropbox client initialization.
        """
        try:
            # Connect signals for token received and auth failed BEFORE initialization
            self.db_initializer.token_received.connect(self.on_token_received)
            self.db_initializer.auth_failed.connect(self.on_auth_failed)

            # Initialize Dropbox client
            self.db_initializer.initialize_dropbox_client()

        except Exception as e:
            logging.error(f"InitializationWorker encountered an error: {e}")
            self.initialization_failed.emit(str(e))

    def on_token_received(self, token):
        """
        Slot triggered when the Dropbox access token is received.
        """
        try:
            # Encrypt and save the token
            with Path("encryption_key.key").open("rb") as f:
                key = f.read()
            fernet = Fernet(key)
            tokens = {"access_token": token}
            encrypted_data = fernet.encrypt(json.dumps(tokens).encode('utf-8'))
            with Path("dropbox_token.json").open("wb") as f:
                f.write(encrypted_data)
            logging.info("Access token encrypted and saved to 'dropbox_token.json'.")
            self.initialization_complete.emit()
        except Exception as e:
            logging.error(f"Failed to encrypt and save access token: {e}")
            self.initialization_failed.emit(f"Failed to encrypt and save access token: {e}")

    def on_auth_failed(self, error):
        """
        Slot triggered when Dropbox authentication fails.
        """
        logging.error(f"Authentication failed: {error}")
        self.initialization_failed.emit(f"Authentication failed: {error}")

# ------------------------------ PyQt5 GUI Class ------------------------------ #

class ClientIntakeWindow(QMainWindow):
    """
    The main window of the Client Intake System application.
    """

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Client Intake System")
        self.setGeometry(100, 100, 500, 700)

        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)
        layout = QFormLayout(central_widget)

        # First Name Input
        self.first_name_input = QLineEdit()
        self.first_name_input.textChanged.connect(self.capitalize_first_letter)
        layout.addRow("Client First Name:", self.first_name_input)

        # Last Name Input
        self.last_name_input = QLineEdit()
        self.last_name_input.textChanged.connect(self.capitalize_first_letter)
        layout.addRow("Client Last Name:", self.last_name_input)

        # Opposing Party Input
        self.opposing_party_input = QLineEdit()
        layout.addRow("Opposing Party:", self.opposing_party_input)

        # Email Input
        self.email_input = QLineEdit()
        layout.addRow("Client Email:", self.email_input)

        # Address Input
        self.address_input = QLineEdit()
        layout.addRow("Client Address:", self.address_input)

        # Phone Number Input
        self.phone_input = QLineEdit()
        layout.addRow("Client Phone:", self.phone_input)

        # Fee Arrangement Dropdown
        self.fee_arrangement_dropdown = QComboBox()
        self.fee_arrangement_dropdown.addItems(["Select Fee Arrangement", "Hourly", "Contingency"])
        self.fee_arrangement_dropdown.currentIndexChanged.connect(self.fee_arrangement_changed)
        layout.addRow("Fee Arrangement:", self.fee_arrangement_dropdown)

        # Hourly Rate Input
        self.hourly_rate_input = QLineEdit()
        self.hourly_rate_input.setPlaceholderText("Enter hourly rate")
        self.hourly_rate_input.setVisible(False)
        layout.addRow("Hourly Rate:", self.hourly_rate_input)

        # Contingency Percentage Input
        self.contingency_percentage_input = QLineEdit()
        self.contingency_percentage_input.setPlaceholderText("Enter contingency percentage (0-100)")
        self.contingency_percentage_input.setVisible(False)
        layout.addRow("Contingency Percentage:", self.contingency_percentage_input)

        # Primary Attorney Input
        self.primary_attorney_input = QLineEdit()
        layout.addRow("Primary Attorney:", self.primary_attorney_input)

        # Is Bosch Checkbox
        self.bosch_checkbox = QCheckBox("Is this a Bosch file?")
        layout.addRow(self.bosch_checkbox)

        # Status Label
        self.status_label = QLabel("Initializing Dropbox and database connection...")
        layout.addRow(self.status_label)

        # Submit Button
        self.submit_button = QPushButton("Submit")
        self.submit_button.clicked.connect(self.on_submit)
        self.submit_button.setEnabled(False)  # Initially disabled until initialization completes
        layout.addRow(self.submit_button)

        # Initialize Dropbox Client
        self.db_initializer = DropboxClientInitializer()

        # Setup and start the initialization thread
        self.init_thread = QThread()
        self.init_worker = InitializationWorker(self.db_initializer)
        self.init_worker.moveToThread(self.init_thread)
        self.init_thread.started.connect(self.init_worker.run)
        self.init_worker.initialization_complete.connect(self.on_initialization_complete)
        self.init_worker.initialization_failed.connect(self.on_initialization_failed)
        self.init_thread.start()

        # Initialize database connection once Dropbox is ready
        self.db_conn = None  # Will be set after Dropbox initialization

        # ------------------------------ Input Validators ------------------------------ #
        from PyQt5.QtGui import QDoubleValidator

        # Set validator for Hourly Rate
        double_validator = QDoubleValidator(0.0, 10000.0, 2)
        self.hourly_rate_input.setValidator(double_validator)

        # Set validator for Contingency Percentage
        double_validator_percentage = QDoubleValidator(0.0, 100.0, 2)
        self.contingency_percentage_input.setValidator(double_validator_percentage)

    def capitalize_first_letter(self):
        """
        Capitalizes the first letter of the input field.
        """
        sender = self.sender()
        text = sender.text()
        if text:
            capitalized_text = text[0].upper() + text[1:]
            if text != capitalized_text:
                sender.blockSignals(True)
                sender.setText(capitalized_text)
                sender.blockSignals(False)

    def fee_arrangement_changed(self, index):
        """
        Shows or hides fee arrangement fields based on the selected option.
        """
        arrangement = self.fee_arrangement_dropdown.currentText()
        if arrangement == "Hourly":
            self.hourly_rate_input.setVisible(True)
            self.contingency_percentage_input.setVisible(False)
        elif arrangement == "Contingency":
            self.hourly_rate_input.setVisible(False)
            self.contingency_percentage_input.setVisible(True)
        else:
            self.hourly_rate_input.setVisible(False)
            self.contingency_percentage_input.setVisible(False)

    def on_initialization_complete(self):
        """
        Slot triggered when initialization is complete.
        Initializes the database connection and enables the Submit button.
        """
        try:
            # Initialize database connection
            self.db_conn = initialize_database(self.db_initializer.db_client, Path(LOCAL_DB_PATH))
            logging.info("Database initialized successfully.")

            # Verify database schema
            verify_database_schema(self.db_conn)

            self.status_label.setText("Initialization complete. You can now submit client data.")
            self.submit_button.setEnabled(True)
            logging.debug("Submit button enabled.")
        except Exception as e:
            logging.error(f"Failed to initialize database after Dropbox authentication: {e}")
            QMessageBox.critical(self, "Database Error", f"Failed to initialize database: {e}")
            self.status_label.setText("Initialization failed. Please check the logs for details.")

    def on_initialization_failed(self, error):
        """
        Slot triggered when initialization fails.
        Displays an error message and keeps the Submit button disabled.
        """
        logging.error(f"Initialization failed: {error}")
        QMessageBox.critical(self, "Initialization Error", f"Failed to initialize application: {error}")
        self.status_label.setText("Initialization failed. Please check the logs for details.")

    def on_submit(self):
        """
        Handles the Submit button click event.
        Validates input, processes client data, and provides feedback to the user.
        """
        client_data = {
            "Client_FirstName": self.first_name_input.text().strip(),
            "Client_LastName": self.last_name_input.text().strip(),
            "Opposing_Party": self.opposing_party_input.text().strip(),
            "Client_Email": self.email_input.text().strip(),
            "Client_Address": self.address_input.text().strip(),
            "Client_Phone": self.phone_input.text().strip(),
            "Fee_Arrangement": self.fee_arrangement_dropdown.currentText() if self.fee_arrangement_dropdown.currentText() in [
                "Hourly", "Contingency"] else "",
            "Hourly_Rate": self.hourly_rate_input.text().strip() if self.fee_arrangement_dropdown.currentText() == "Hourly" else None,
            "Contingency_Percentage": self.contingency_percentage_input.text().strip() if self.fee_arrangement_dropdown.currentText() == "Contingency" else None,
            "Primary_Attorney": self.primary_attorney_input.text().strip(),
            "IsBosch": "True" if self.bosch_checkbox.isChecked() else "False"
        }

        logging.info("Submit button clicked. Processing client data.")

        try:
            if not hasattr(self, 'db_conn') or not self.db_conn:
                error_msg = "Database connection is not initialized."
                logging.error(error_msg)
                QMessageBox.critical(self, "Database Error", error_msg)
                return

            # Additional Validations
            if client_data['Fee_Arrangement'] == "Hourly":
                if not client_data['Hourly_Rate']:
                    raise ValueError("Hourly Rate is required for Hourly fee arrangement.")
                if not re.match(r'^\d+(\.\d{1,2})?$', client_data['Hourly_Rate']):
                    raise ValueError("Hourly Rate must be a valid number (up to two decimal places).")
            elif client_data['Fee_Arrangement'] == "Contingency":
                if not client_data['Contingency_Percentage']:
                    raise ValueError("Contingency Percentage is required for Contingency fee arrangement.")
                if not re.match(r'^\d+(\.\d{1,2})?$', client_data['Contingency_Percentage']):
                    raise ValueError("Contingency Percentage must be a valid number (up to two decimal places).")

            # Process client data
            client_name = process_client_data(self.db_initializer.db_client, self.db_conn, client_data)
            QMessageBox.information(self, "Success", f"Client intake process completed successfully for {client_name}")
            logging.info(f"Client intake process completed successfully for {client_name}.")

            # Clear the form after successful submission
            self.clear_form()

        except ValueError as ve:
            QMessageBox.warning(self, "Validation Error", str(ve))
            logging.warning(f"Validation error: {ve}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {e}")
            logging.error(f"An unexpected error occurred: {e}", exc_info=True)

    def clear_form(self):
        """
        Clears all input fields after successful submission.
        """
        self.first_name_input.clear()
        self.last_name_input.clear()
        self.opposing_party_input.clear()
        self.email_input.clear()
        self.address_input.clear()
        self.phone_input.clear()
        self.fee_arrangement_dropdown.setCurrentIndex(0)
        self.hourly_rate_input.clear()
        self.contingency_percentage_input.clear()
        self.primary_attorney_input.clear()
        self.bosch_checkbox.setChecked(False)

# ------------------------------ Helper Functions ------------------------------ #

def verify_database_schema(conn):
    """
    Verifies that all required columns exist in the 'Clients' table.

    Args:
        conn (pyodbc.Connection): Active database connection.

    Raises:
        ValueError: If any required columns are missing.
    """
    table_name = "Clients"
    expected_columns = [
        "Client_FirstName", "Client_LastName", "Opposing_Party", "Client_Email",
        "Client_Address", "Client_Phone", "Fee_Arrangement", "Contingency_Amount",
        "Hourly_Amount", "Primary_Attorney", "EncryptedData", "IsBosch"
    ]

    actual_columns = list_table_columns(conn, table_name)
    missing_columns = [col for col in expected_columns if col not in actual_columns]

    if missing_columns:
        error_msg = f"Missing columns in '{table_name}' table: {missing_columns}"
        logging.error(error_msg)
        raise ValueError(error_msg)
    else:
        logging.info(f"All required columns are present in the '{table_name}' table.")

# ------------------------------ Main Application Entry Point ------------------------------ #

def main():
    """
    The main entry point for the Client Intake System application.
    """
    app = QApplication(sys.argv)
    window = ClientIntakeWindow()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
